from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path("/Users/wangyujie/Desktop/Career Aspirations and Strategic Plan.docx")
TEMPLATE_PATH = Path("/Users/wangyujie/Desktop/Career Plan and Skills Reflection Template.docx")
WORKSPACE_OUTPUT = Path("/Users/wangyujie/Desktop/web_lin/Career Plan and Skills Reflection Template_filled.docx")


AMBITION = (
    "My long-term ambition is to build an international career as an opera accompanist, "
    "repetiteur and collaborative pianist, while also working in music education. I hope to "
    "support singers through performance and coaching, keep developing as a teacher, and use "
    "my work to connect Chinese and Western musical cultures in a meaningful way."
)

SKILLS_REQUIRED = [
    "Advanced collaborative piano technique and secure score preparation for opera, art song and choral repertoire",
    "Vocal coaching skills, including language awareness, diction support, breathing sensitivity and stylistic understanding",
    "Efficient rehearsal-room communication with singers, conductors, directors and instrumental colleagues",
    "Teaching ability in piano, vocal coaching, musicianship, theory and examination preparation",
    "Professional resilience in auditions, networking, self-management and portfolio career building",
    "Evidence of academic capability for conservatoire or higher-education teaching roles, including research and writing",
    "Project management, concert organisation and the ability to build reliable artistic partnerships",
    "Digital branding, audience engagement and the confidence to present work consistently across platforms",
]

SKILLS_ALIGNMENT = (
    "Looking across the five opportunities in my portfolio, these are the skills that appear most "
    "often. I feel that I already match well with the areas of collaborative piano, vocal "
    "understanding, teaching experience and cross-cultural communication. The main areas I still "
    "need to strengthen are higher-level opera experience, a stronger academic profile, more "
    "consistent self-promotion, and clearer evidence of interdisciplinary work."
)

STRENGTHS_PARAGRAPHS = [
    (
        "One of my biggest strengths at the moment is the combination of collaborative piano "
        "training, vocal knowledge and teaching experience. I do not think about accompaniment as "
        "just playing the notes accurately. Because I have trained in both piano and vocal "
        "performance, I naturally listen from the singer's side as well. I pay attention to breath, "
        "language, phrasing and balance, and this helps me work in a more supportive and musical way. "
        "In rehearsals, I often find that this background helps me notice problems quickly and respond "
        "in a practical way. For example, I can often tell whether a difficulty comes from breathing, "
        "diction, pacing or confidence rather than from the music alone. I think this is one of the "
        "main reasons why collaborative work suits me so well."
    ),
    (
        "My postgraduate study in the UK has also helped me grow a lot. At the Royal College of "
        "Music, I had the chance to work more deeply on vocal repertoire and became more confident "
        "with Italian, German and French music. This improved not only my language awareness, but "
        "also my stylistic understanding and the way I support singers in performance. Later, at the "
        "Royal Birmingham Conservatoire, my Artist Diploma focused more directly on opera production. "
        "That experience helped me understand the professional side of opera more clearly, especially "
        "how rehearsals are structured and how different artistic decisions are made. Working with "
        "people such as Mark Austin and Lesanne Van Overbeek was especially valuable because it gave "
        "me a clearer picture of the level expected in the profession."
    ),
    (
        "Another important strength is my educational background. During my undergraduate degree in "
        "music education, I studied pedagogy, psychology, and both Western and Chinese music history. "
        "This gave me a strong foundation for teaching, and it also taught me to think carefully "
        "about how students learn rather than only what I want to teach. I also hold a Chinese music "
        "teaching qualification and have learnt about methods such as Suzuki, Kodaly and Dalcroze. "
        "Because of this, I feel comfortable adapting my teaching to different students and different "
        "ages. So far, I have taught students from the ages of six to twenty-five in piano, vocal "
        "coaching, music theory and ABRSM preparation, which has made me much more flexible and "
        "practical as a teacher."
    ),
    (
        "My teaching experience has also shown me that I can work with both individuals and groups. "
        "During my six-month teaching placement at a Chinese high school, I learnt how to manage a "
        "classroom, rehearse choirs and support students over a longer period of time. This was a very "
        "useful experience because it pushed me to become more organised, patient and clear in my "
        "communication. It also showed me that good teaching is not only about musical knowledge, but "
        "also about responsibility, planning and understanding students' needs. I think this is an "
        "important part of my profile because many of the roles I am interested in combine artistry "
        "with teaching and mentoring."
    ),
    (
        "Performance experience is another area where I feel I already have a strong base. My work so "
        "far has included different kinds of performances, from projects linked with Birmingham Opera "
        "Company to lunchtime solo recitals in London churches. These experiences have helped me become "
        "more adaptable, because each setting requires a different kind of preparation, communication "
        "and stage awareness. I have also become a signed artist with Hemiola Music Company, which has "
        "helped me begin building a more public professional identity. Although I still see myself as "
        "developing, I do feel that I already have real performing experience rather than only academic "
        "training."
    ),
    (
        "Finally, I think one of my most valuable strengths is my cross-cultural perspective. My "
        "training and experience already connect Chinese and British musical environments, and I feel "
        "this has shaped the way I see my future career. I am interested not only in performing and "
        "teaching, but also in helping different musical cultures speak to each other. This is why I "
        "am especially drawn to repertoire, projects and teaching opportunities that involve cultural "
        "exchange. Overall, I feel that my strongest qualities at present are musicianship, sensitivity "
        "to singers, teaching ability and an international outlook. These give me a solid foundation "
        "for the kind of career I want to build."
    ),
]

IMPROVEMENT_PARAGRAPHS = [
    (
        "When I compare myself with the criteria in the five professional opportunities, the area that "
        "stands out most is high-level professional placement. I feel that I already have solid "
        "training and relevant experience, but for very competitive roles such as young artist "
        "programmes, opera house repetiteur positions or top conservatoire jobs, employers often look "
        "for stronger evidence of professional recognition. Reaching the final round of the National "
        "Opera Studio selection for 2026/27 was encouraging for me because it showed that I am already "
        "close to that level. At the same time, it also reminded me that I still need to improve my "
        "audition preparation, confidence under pressure and overall professional polish."
    ),
    (
        "A closely related area is repertoire depth and audition strategy. I already have a good "
        "foundation in vocal repertoire and language-based work, but I know that professional auditions "
        "require more than being generally musical. They require very well-prepared excerpts, quick "
        "stylistic awareness, strong sight-reading and the ability to stay calm and convincing in a "
        "short space of time. At the moment, I think I can compete, but I would like to move beyond "
        "being competitive and become much more secure and consistent. For that reason, I want to keep "
        "developing a more focused audition repertoire and get more experience in mock auditions and "
        "professional coaching."
    ),
    (
        "For teaching-related roles, I feel that my main weakness at present is my academic profile. I "
        "have a good background in pedagogy and useful teaching experience, but many conservatoire or "
        "university posts also expect publications, research activity or conference participation. This "
        "is something I still need to build more seriously. At the moment, I would describe myself as "
        "someone with strong practical teaching ability, but in the future I would like to become a "
        "more complete practitioner-educator whose performance work and academic work support each other."
    ),
    (
        "Another area that needs improvement is the way I present myself professionally. I already have "
        "real performance experience and useful contacts, but I know that my digital presence is not yet "
        "as strong or as organised as it should be. I need a clearer press kit, better online material, "
        "more regular social-media updates and a stronger sense of how I want to introduce myself as both "
        "a collaborative pianist and a teacher. I used to think this was a secondary issue, but I now "
        "realise it matters because it affects how opportunities come to me and how other people remember "
        "my work."
    ),
    (
        "I also feel that I need to build a more stable network of long-term collaborators. I already "
        "have some valuable connections with singers, teachers and institutions, but many of these are "
        "still project-based rather than long term. In the future, I would like to have more regular "
        "artistic partners and mentors, because this would help me build stronger ensemble work, a more "
        "recognisable profile and more consistent opportunities. I think this is especially important in "
        "the music profession, where trust and long-term collaboration often lead to the best work."
    ),
    (
        "Finally, I am very interested in interdisciplinary work, especially music therapy and the "
        "promotion of Chinese art song, but I know these interests are still at an early stage. Right "
        "now, they are more like important directions that I want to explore further rather than fully "
        "developed parts of my career. If I want them to become serious professional strengths, I will "
        "need more training, more project experience and clearer evidence of the impact they can make. "
        "For me, this is an exciting area for future development, because it could help me build a career "
        "that is not only artistic, but also socially meaningful and distinctive."
    ),
]

ACTION_PLAN_ROWS = [
    (
        "Audition repertoire, opera coaching profile and professional placement",
        "Build a yearly audition calendar; refine core opera excerpts with mock auditions; seek regular coaching from conductors, repetiteurs and vocal staff; apply strategically to young artist programmes, opera studios and accompanist posts.",
        "Short term: within 12 months; ongoing review over 5 years",
    ),
    (
        "Work eligibility and international employability",
        "Secure the PSW visa and maintain up-to-date right-to-work documentation; monitor UK and international opportunities; prepare application materials so that I can respond quickly when freelance or salaried roles arise.",
        "Within 1 year",
    ),
    (
        "Academic profile, research writing and conservatoire readiness",
        "Identify research topics connected to vocal collaboration, pedagogy or Chinese art song; submit articles to peer-reviewed journals; attend academic conferences and build evidence of reflective teaching practice.",
        "Medium term: 2 to 5 years",
    ),
    (
        "Digital branding, press materials and audience development",
        "Create a professional website press kit, biography, repertoire list and performance clips; post consistently on Instagram, YouTube, Threads, Xiaohongshu, Weibo and WeChat; turn performances and rehearsals into portfolio content.",
        "Foundation within 1 year; strong platform within 3 to 5 years",
    ),
    (
        "Interdisciplinary and cross-cultural project development",
        "Develop alumni concerts, Chinese art song advocacy projects and pilot music-therapy collaborations; document outcomes, gather feedback and use these projects to demonstrate leadership, cultural engagement and social impact.",
        "Pilot projects within 1 to 2 years; mature profile within 5 to 10 years",
    ),
]

EXTERNAL_PROJECT = (
    "For the external engagement project, I would choose to develop and present an alumni concert "
    "series with soprano graduates from Zhejiang Conservatory of Music and Shenzhen Art School. I think "
    "this is a very suitable project for me because it brings together the main parts of my career plan: "
    "collaborative performance, networking and cultural exchange. As the pianist and one of the organisers, "
    "I would not only prepare the music, but also help with communication, programming and promotion. "
    "Because of this, the project would feel like a real professional step rather than only an academic task."
)

EXTERNAL_PROJECT_2 = (
    "The value of this project is that it would directly strengthen several areas that I want to improve. "
    "First, it would give me more collaborative experience with singers in a public performance setting, "
    "which is very relevant to my goal of working as an opera accompanist and vocal coach. Second, it would "
    "help me reconnect with musicians, institutions and possible employers in China, especially if the "
    "concerts were supported by local media coverage such as Shenzhen Daily. Third, it would give me useful "
    "material for my professional portfolio, including performance recordings, publicity material and clear "
    "evidence that I can initiate and organise projects."
)

EXTERNAL_PROJECT_3 = (
    "Most importantly, this project reflects the kind of musician I hope to become. I do not want to be "
    "only a pianist who waits for opportunities. I want to be someone who can create opportunities, support "
    "other artists and help connect different musical communities. In that sense, this project would not sit "
    "separately from my career plan. It would be a very practical way of moving that plan forward."
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_paragraph.style = style
        except Exception:
            pass
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def paragraph_with_answer(paragraph: Paragraph, answer: str) -> Paragraph:
    answer_paragraph = insert_paragraph_after(paragraph, answer)
    return answer_paragraph


def find_paragraph(doc: Document, expected_text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == expected_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {expected_text}")


def find_paragraph_startswith(doc: Document, expected_prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(expected_prefix):
            return paragraph
    raise ValueError(f"Paragraph not found with prefix: {expected_prefix}")


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_between(doc: Document, start_text: str, end_text: str) -> None:
    start = find_paragraph(doc, start_text)
    end = find_paragraph(doc, end_text)
    start_element = start._element
    end_element = end._element
    in_range = False
    to_remove: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        if paragraph._element == start_element:
            in_range = True
            continue
        if paragraph._element == end_element:
            break
        if in_range and paragraph.text.strip():
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def clear_after(doc: Document, start_text: str) -> None:
    start = find_paragraph(doc, start_text)
    start_element = start._element
    seen_start = False
    to_remove: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        if paragraph._element == start_element:
            seen_start = True
            continue
        if seen_start and paragraph.text.strip():
            to_remove.append(paragraph)
    for paragraph in to_remove:
        remove_paragraph(paragraph)


def main() -> None:
    doc = Document(TEMPLATE_PATH)
    name_paragraph = find_paragraph_startswith(doc, "Name:")
    name_paragraph.text = "Name: Yizhe Lin"

    ambition_text = "What is your ultimate career ambition(s)? (approx. 50 words)"
    reflect_text = (
        "Reflect on the five job advertisements that you have included in the portfolio. What skills are required for each and how do they align with your own skills at present? Where are your strengths and where are areas that you feel could be improved?"
    )
    skills_text = "Skills required (this can be a bullet point list)"
    strengths_text = (
        "What are your present strengths? What are you good at? What relevant experience have you already acquired (approximately 750 words)"
    )
    improvement_text = (
        "Reflecting on the five professional career opportunities and their criteria, which required skills/experience do you feel could be improved and why? Where do you feel your skills/experience are at present and where would you like them to be? (approx. 750 words)"
    )
    action_text = (
        "Now create a career action plan in order to help you achieve those goals. Think about short-term (1 year), medium term (5 years) and long term (10 years). What further experience/training/skill development do you require and how can you achieve this? This table can be expanded according to the number of skills discussed"
    )
    rationale_text = (
        "Please explain which option you are choosing, describe your chosen project and explain how this will assist you with the career plan discussed above (approx. 250 words)"
    )

    clear_between(doc, ambition_text, reflect_text)
    clear_between(doc, skills_text, strengths_text)
    clear_between(doc, strengths_text, improvement_text)
    clear_between(doc, improvement_text, action_text)
    clear_after(doc, rationale_text)

    ambition_prompt = find_paragraph(doc, ambition_text)
    current = paragraph_with_answer(ambition_prompt, AMBITION)

    skills_prompt = find_paragraph(doc, skills_text)
    current = insert_paragraph_after(skills_prompt, "")
    for item in SKILLS_REQUIRED:
        current = insert_paragraph_after(current, f"- {item}")
    current = insert_paragraph_after(current, SKILLS_ALIGNMENT)

    strengths_prompt = find_paragraph(doc, strengths_text)
    current = paragraph_with_answer(strengths_prompt, STRENGTHS_PARAGRAPHS[0])
    for text in STRENGTHS_PARAGRAPHS[1:]:
        current = insert_paragraph_after(current, text)

    improvement_prompt = find_paragraph(doc, improvement_text)
    current = paragraph_with_answer(improvement_prompt, IMPROVEMENT_PARAGRAPHS[0])
    for text in IMPROVEMENT_PARAGRAPHS[1:]:
        current = insert_paragraph_after(current, text)

    table = doc.tables[0]
    for row_data, row in zip(ACTION_PLAN_ROWS, table.rows[1:]):
        for value, cell in zip(row_data, row.cells):
            cell.text = value

    project_prompt = find_paragraph(doc, rationale_text)
    current = paragraph_with_answer(project_prompt, EXTERNAL_PROJECT)
    current = insert_paragraph_after(current, EXTERNAL_PROJECT_2)
    insert_paragraph_after(current, EXTERNAL_PROJECT_3)

    doc.save(WORKSPACE_OUTPUT)


if __name__ == "__main__":
    main()
